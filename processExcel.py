import json
import os
import logging

import xlrd3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

"""
    prepare: pip3 install xlrd3 python-docx pyinstaller
    package: pyinstaller -F .\processExcel.py 
             cp .\dist\processExcel.exe .  
"""

# 配置日志记录级别和输出格式
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 创建日志记录器
logger = logging.getLogger(__name__)

# 创建文件处理器，用于保存错误日志到文件


file_handler = logging.FileHandler('error.log')
file_handler.setLevel(logging.ERROR)  # 设置处理器的日志记录级别

# 创建控制台处理器，用于将其他级别的日志打印到标准输出
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)  # 设置处理器的日志记录级别

# 配置处理器的输出格式
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

# 将处理器添加到日志记录器
logger.addHandler(file_handler)
logger.addHandler(console_handler)


excel_header_info = dict()
conf = {}

"""
    预处理每一行数据，符合填充要求，然后填充
"""
def process_line(sheet, row):
    global conf

    DELIVERY_UNIT_IDX = conf.get("DELIVERY_UNIT_IDX")
    DELIVERY_ADDRESS_IDX = conf.get("DELIVERY_ADDRESS_IDX")
    DELIVERY_ADDRESS_PERSON_IDX = conf.get("DELIVERY_ADDRESS_PERSON_IDX")
    DELIVERY_ADDRESS_PHONE_IDX = conf.get("DELIVERY_ADDRESS_PHONE_IDX")

    INT_POSITION = conf.get("INT_POSITION")

    logger.info("process line {}".format(row))
    line = []
    for col in range(sheet.ncols):
        # 获取单元格的值
        value = sheet.cell_value(row, col)

        if isinstance(value, str):
            value = value.strip()
            value = value.replace(' ', '')

        # 序号，防止变成浮点数
        if col in INT_POSITION:
            if isinstance(value, str) and len(value) == 0:
                value = 0
            else:
                value = int(float(value))
        if col == DELIVERY_UNIT_IDX:
            value = '收货单位：' + str(value)
        if col == DELIVERY_ADDRESS_IDX:
            value = '收货地址：' + str(value)
        # 拼接收件人+ 地址
        if col == DELIVERY_ADDRESS_PERSON_IDX:
            phone = str(sheet.cell_value(row, DELIVERY_ADDRESS_PHONE_IDX)).strip().replace(' ', '')
            value = '收货联系人：' + value + ' ' + str(int(float(phone)))

        line.append(value)
    process_docx(line)


def process_file(input_file):
    INPUT_COLUMN_MAPPING = conf.get("INPUT_COLUMN_MAPPING")

    with xlrd3.open_workbook(input_file) as wb:
        for sheet in wb.sheets():
            logger.info("开始处理文件[{}-{}]".format(input_file, sheet.name))
            if sheet.ncols < len(INPUT_COLUMN_MAPPING):
                logger.error("文件{}-{}列数少于预期，跳过处理".format(input_file, sheet.name))
            # 获取 header 信息
            for col in range(sheet.ncols):
                key = INPUT_COLUMN_MAPPING[col]
                value = sheet.cell_value(0, col)
                excel_header_info[key] = value
            for row in range(1, sheet.nrows):
                process_line(sheet, row)
            break


def process():
    check_output_path()
    for dirpath, dirnames, filenames in os.walk('./input'):
        for filename in filenames:
            if filename.find("~$") != -1:
                continue
            process_file(dirpath + '/' + filename)


def check_output_path():
    OUTPUT_PATH = conf.get("OUTPUT_PATH")

    if not os.path.exists(OUTPUT_PATH):
        os.makedirs(OUTPUT_PATH)
    else:
        files = os.listdir(OUTPUT_PATH)
        # 遍历所有的文件和文件夹
        for file in files:
            # 拼接完整的路径
            file_path = os.path.join(OUTPUT_PATH, file)
            # 判断是否是文件
            if os.path.isfile(file_path):
                # 删除文件
                os.remove(file_path)


def process_docx(line):
    TEMPLATE_PATH = conf.get("TEMPLATE_PATH")
    title_text = conf.get("TITLE")
    HEADER_POSITION_MAPPING = conf.get("HEADER_POSITION_MAPPING")
    OUTPUT_POSITION_MAPPING = conf.get("OUTPUT_POSITION_MAPPING")
    INPUT_COLUMN_MAPPING = conf.get("INPUT_COLUMN_MAPPING")
    addtion_infos = conf.get("ADDTIONAL_INFO")
    OUTPUT_PATH = conf.get("OUTPUT_PATH")
    BOLD_POSITION = conf.get("BOLD_POSITION")
    CENTERED_POSITION = conf.get("CENTERED_POSITION")


    document = Document(TEMPLATE_PATH)

    ## 处理标题  hrad code
    title_paragh = document.paragraphs[1]
    # 设置段落格式为黑体小二居中
    title_paragh.text = title_text
    title_paragh.runs[0].font.name = "Arial"
    title_paragh.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    title_paragh.runs[0].font.size = Pt(18)
    title_paragh.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ## 处理表格部分
    table = document.tables[0]

    ## 处理表头部分
    for key in HEADER_POSITION_MAPPING:
        pos = HEADER_POSITION_MAPPING.get(key)
        value = excel_header_info.get(key)
        cell = table.cell(pos[0], pos[1])
        cell.text = str(value)

    # 补充excel中每一行的信息
    for idx in range(len(line)):
        pos = OUTPUT_POSITION_MAPPING.get(INPUT_COLUMN_MAPPING[idx])
        if pos is None:
            continue
        cell = table.cell(pos[0], pos[1])
        cell.text = str(line[idx])

        
    # 补充额外信息
    for key in addtion_infos:
        pos = OUTPUT_POSITION_MAPPING.get(key)
        if pos is None:
            continue
        cell = table.cell(pos[0], pos[1])
        cell.text = addtion_infos.get(key)

    # 设置中心对称  & 加粗
    for pos in BOLD_POSITION:
        cell = table.cell(pos[0], pos[1])
        cell.paragraphs[0].runs[0].bold = True

    for pos in CENTERED_POSITION:
        cell = table.cell(pos[0], pos[1])
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # document.add_picture("./1.jpg")  # 等同于doc.add_paragraph().add_run().add_picture()
    # document.save('./test.docx')

    output_file_name = OUTPUT_PATH + '/' + '{}、{}.docx'.format(line[0], line[1])
    document.save(output_file_name)


def load_conf():
    global conf

    with open('./setting.json', 'r', encoding='utf-8') as f:
        conf = json.load(f)

        INPUT_COLUMN_MAPPING = conf.get("INPUT_COLUMN_MAPPING")
        OUTPUT_POSITION_MAPPING = conf.get("OUTPUT_POSITION_MAPPING")
        HEADER_POSITION_MAPPING = conf.get("HEADER_POSITION_MAPPING")
        OUTPUT_PATH = conf.get("OUTPUT_PATH")
        TEMPLATE_PATH = conf.get("TEMPLATE_PATH")
        DELIVERY_ADDRESS_PHONE_IDX = int(conf.get("DELIVERY_ADDRESS_PHONE_IDX"))
        DELIVERY_UNIT_IDX = int(conf.get("DELIVERY_UNIT_IDX"))
        DELIVERY_ADDRESS_PERSON_IDX = int(conf.get("DELIVERY_ADDRESS_PERSON_IDX"))
        DELIVERY_ADDRESS_IDX = int(conf.get("DELIVERY_ADDRESS_IDX"))
        

        assert(DELIVERY_ADDRESS_PHONE_IDX != -1)
        assert(DELIVERY_UNIT_IDX != -1)
        assert(DELIVERY_ADDRESS_PERSON_IDX != -1)
        assert(DELIVERY_ADDRESS_IDX != -1)

        addtion_infos = conf.get("ADDTIONAL_INFO")
        title_text = conf.get("TITLE")
    pass

def main():
    load_conf()
    process()


if __name__ == "__main__":
    main()
